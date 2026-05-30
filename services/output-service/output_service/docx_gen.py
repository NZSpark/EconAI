"""DOCX generator compliant with GB/T 9704-2012 (M7-09 through M7-16).

Generates Chinese government document format .docx files using python-docx.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from output_service.template_loader import get_template_loader

REF_PATTERN = re.compile(r"\[ref:([^\]]+)\]")

# 回退 font mapping when fonts are unavailable
_FONT_FALLBACKS = {
    "小标宋_GB2312": "宋体",
    "仿宋_GB2312": "仿宋",
    "楷体_GB2312": "楷体",
    "黑体": "黑体",
}


def _set_font(run: Any, font_name: str, size_pt: float, bold: bool = False, color: str | None = None) -> None:
    """Set font properties on a run."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def _add_paragraph(
    doc: Any,  # python-docx Document
    text: str,
    font_name: str,
    size_pt: float,
    alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent_chars: int = 0,
    line_spacing_pt: float | None = None,
    line_spacing: float | None = None,
    bold: bool = False,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> Any:
    """Add a styled paragraph to the document."""
    para = doc.add_paragraph()
    para.alignment = alignment
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)

    if first_line_indent_chars:
        pf.first_line_indent = Cm(first_line_indent_chars * 0.74)

    if line_spacing is not None:
        pf.line_spacing = line_spacing
    elif line_spacing_pt is not None:
        pf.line_spacing = Pt(line_spacing_pt)

    run = para.add_run(text)
    _set_font(run, font_name, size_pt, bold)
    return para


def _build_footnote_numbering(text: str, ref_map: dict[str, int]) -> tuple[str, list[str]]:
    """Replace [ref:xxx] with superscript [n] markers, return text and ordered refs."""
    ordered_refs: list[str] = []

    def _replacer(match: re.Match[str]) -> str:
        ref_id = match.group(1)
        if ref_id not in ref_map:
            ref_map[ref_id] = len(ref_map) + 1
            ordered_refs.append(ref_id)
        return f"[{ref_map[ref_id]}]"

    result = REF_PATTERN.sub(_replacer, text)
    return result, ordered_refs


class DocxGenerator:
    """Generates GB/T 9704 compliant .docx files."""

    def __init__(self, institution_name: str | None = None) -> None:
        self._institution_name = institution_name or "PolicyAI 分析中心"
        self._template = get_template_loader().load_docx_template()

    def generate(
        self,
        title: str,
        sections: list[dict[str, Any]],
        citations: list[dict[str, Any]],
        metadata: dict[str, Any] | None = None,
    ) -> bytes:
        """生成 a GB/T 9704 formatted .docx file.

        Args:
            title: Document title.
            sections: List of {"title": str, "level": int, "content": str} dicts.
            citations: List of citation dicts.
            metadata: Optional metadata.

        Returns:
            DOCX file bytes.
        """
        meta = metadata or {}
        doc = Document()
        cfg = self._template

        # --- Page setup ---
        section = doc.sections[0]
        margins = cfg.page_margins
        section.top_margin = Cm(margins.get("top_mm", 37) / 10)
        section.bottom_margin = Cm(margins.get("bottom_mm", 35) / 10)
        section.left_margin = Cm(margins.get("left_mm", 28) / 10)
        section.right_margin = Cm(margins.get("right_mm", 26) / 10)

        # --- Header (版头) ---
        header_cfg = cfg.header
        if header_cfg.get("enabled", True):
            header = doc.sections[0].header
            header.is_linked_to_previous = False
            h_content = header_cfg.get("content", {})
            inst_font = h_content.get("institution_name_font", {})
            h_para = header.paragraphs[0]
            h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_name = meta.get("institution_name", self._institution_name)
            run = h_para.add_run(inst_name)
            _set_font(
                run,
                inst_font.get("name", "小标宋_GB2312"),
                inst_font.get("size_pt", 18),
                bold=inst_font.get("bold", False),
                color=inst_font.get("color", "#CC0000"),
            )

            # Separator line
            sep_cfg = h_content.get("separator_line", {})
            if sep_cfg.get("enabled", True):
                sep_para = header.add_paragraph()
                sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sep_run = sep_para.add_run("─" * 40)
                sep_run.font.size = Pt(8)
                sep_run.font.color.rgb = RGBColor.from_string(
                    sep_cfg.get("color", "#CC0000").lstrip("#")
                )

            # Issue number
            issue_num = meta.get("issue_number", "")
            if issue_num:
                in_font = h_content.get("issue_number_font", {})
                hn_para = header.add_paragraph()
                hn_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = hn_para.add_run(issue_num)
                _set_font(run, in_font.get("name", "仿宋_GB2312"), in_font.get("size_pt", 14))

        # --- Title ---
        fonts = cfg.fonts
        title_font = fonts.get("title_h1", {})
        _add_paragraph(
            doc,
            title,
            font_name=title_font.get("name", "小标宋_GB2312"),
            size_pt=title_font.get("size_pt", 22),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_chars=0,
            line_spacing_pt=36,
            space_before_pt=12,
            space_after_pt=12,
        )

        # --- Recipient (主送机关, optional) ---
        recipient = meta.get("recipient")
        if recipient:
            _add_paragraph(
                doc,
                recipient,
                font_name=fonts.get("body", {}).get("name", "仿宋_GB2312"),
                size_pt=16,
                first_line_indent_chars=0,
                line_spacing=1.5,
                space_after_pt=6,
            )

        # --- Body content ---
        ref_map: dict[str, int] = {}
        body_font = fonts.get("body", {})
        body_font_name = body_font.get("name", "仿宋_GB2312")
        body_font_size = body_font.get("size_pt", 16)

        heading_fonts = {
            1: fonts.get("title_h2", {"name": "黑体", "size_pt": 16}),
            2: fonts.get("title_h3", {"name": "楷体_GB2312", "size_pt": 16}),
        }

        for section_data in sections:
            section_title = section_data.get("title", "")
            section_level = section_data.get("level", 1)
            section_content = section_data.get("content", "")

            # Section heading
            if section_title:
                hfont = heading_fonts.get(section_level, heading_fonts[1])
                heading_label = ""
                heading_mapping = cfg.body_structure.get("heading_mapping", {})
                if section_level == 1:
                    heading_label = heading_mapping.get("level_1", "一、")
                elif section_level == 2:
                    heading_label = heading_mapping.get("level_2", "(一)")

                _add_paragraph(
                    doc,
                    f"{heading_label}{section_title}" if heading_label else section_title,
                    font_name=hfont.get("name", "黑体"),
                    size_pt=hfont.get("size_pt", 16),
                    first_line_indent_chars=0,
                    line_spacing=1.5,
                    space_before_pt=6,
                    space_after_pt=6,
                )

            # Section content with citation replacement
            processed_content, _ = _build_footnote_numbering(section_content, ref_map)
            _add_paragraph(
                doc,
                processed_content,
                font_name=body_font_name,
                size_pt=body_font_size,
                first_line_indent_chars=2,
                line_spacing=1.5,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            )

        # --- Attachment note (附件说明, optional) ---
        attachment = meta.get("attachment")
        if attachment:
            _add_paragraph(
                doc,
                f"附件：{attachment}",
                font_name=body_font_name,
                size_pt=body_font_size,
                first_line_indent_chars=2,
                line_spacing=1.5,
                space_before_pt=6,
            )

        # --- Signature (发文机关署名) ---
        sig_name = meta.get("signature", self._institution_name)
        _add_paragraph(
            doc,
            sig_name,
            font_name=body_font_name,
            size_pt=body_font_size,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            line_spacing=1.5,
            space_before_pt=12,
        )

        # --- Date ---
        doc_date = meta.get("date", "")
        _add_paragraph(
            doc,
            doc_date,
            font_name=body_font_name,
            size_pt=body_font_size,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            line_spacing=1.5,
            space_after_pt=6,
        )

        # --- Reference list (参考文献) ---
        ref_cfg = cfg.reference_list
        ref_title_font = ref_cfg.get("title_font", {})
        _add_paragraph(
            doc,
            ref_cfg.get("title", "参考文献"),
            font_name=ref_title_font.get("name", "黑体"),
            size_pt=ref_title_font.get("size_pt", 16),
            line_spacing=1.5,
            space_before_pt=12,
            space_after_pt=6,
        )

        entry_font = ref_cfg.get("entry_font", {})
        entry_format = ref_cfg.get("entry_format", "[{index}] {authors}. {title}. {source}. {year}. {page_range}.")
        entry_spacing = ref_cfg.get("entry_spacing_pt", 22)

        # Build citation lookup
        cit_lookup: dict[str, dict[str, Any]] = {c.get("ref_id", ""): c for c in citations}

        for i, ref_id in enumerate(ref_map, 1):
            cit = cit_lookup.get(ref_id, {})
            entry_text = entry_format.format(
                index=i,
                authors=cit.get("authors", ""),
                title=cit.get("document_title", cit.get("title", ref_id)),
                source=cit.get("source", cit.get("source_page", "")),
                year=cit.get("year", ""),
                page_range=cit.get("page_range", cit.get("source_page", "")),
            )
            _add_paragraph(
                doc,
                entry_text,
                font_name=entry_font.get("name", "仿宋_GB2312"),
                size_pt=entry_font.get("size_pt", 14),
                first_line_indent_chars=0,
                line_spacing_pt=entry_spacing,
            )

        # --- Footer (版记) ---
        footer_cfg = cfg.footer
        if footer_cfg.get("enabled", True):
            footer = doc.sections[0].footer
            footer.is_linked_to_previous = False
            f_content = footer_cfg.get("content", {})

            # Separator
            f_sep = f_content.get("separator_line", {})
            if f_sep.get("enabled", True):
                fsep_para = footer.add_paragraph()
                fsep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fsep_run = fsep_para.add_run("─" * 40)
                fsep_run.font.size = Pt(6)

            # CC list
            cc_list = meta.get("cc_list", "")
            if cc_list:
                cc_font = f_content.get("cc_font", {})
                cc_para = footer.add_paragraph()
                cc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = cc_para.add_run(f"抄送：{cc_list}")
                _set_font(run, cc_font.get("name", "仿宋_GB2312"), cc_font.get("size_pt", 14))

            # Issue date
            issue_date = meta.get("date", "")
            if issue_date:
                id_font = f_content.get("issue_date_font", {})
                id_para = footer.add_paragraph()
                id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = id_para.add_run(f"印发日期：{issue_date}")
                _set_font(run, id_font.get("name", "仿宋_GB2312"), id_font.get("size_pt", 14))

        # Write to bytes
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
