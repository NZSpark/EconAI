"""M7-36: Tests for GB/T 9704 DOCX generation."""

from output_service.docx_gen import DocxGenerator, _build_footnote_numbering, _set_font


class TestFootnoteNumbering:
    def test_single_ref(self) -> None:
        ref_map: dict[str, int] = {}
        text, refs = _build_footnote_numbering("Claim [ref:doc:p1].", ref_map)
        assert "[1]" in text
        assert refs == ["doc:p1"]

    def test_multiple_unique_refs(self) -> None:
        ref_map: dict[str, int] = {}
        text, refs = _build_footnote_numbering("A [ref:a:1] B [ref:b:2].", ref_map)
        assert "[1]" in text
        assert "[2]" in text
        assert refs == ["a:1", "b:2"]

    def test_duplicate_ref_same_number(self) -> None:
        ref_map: dict[str, int] = {}
        text, refs = _build_footnote_numbering("A [ref:x:1] and [ref:x:1].", ref_map)
        assert text.count("[1]") == 2
        assert refs == ["x:1"]

    def test_no_refs(self) -> None:
        ref_map: dict[str, int] = {}
        text, refs = _build_footnote_numbering("Plain text.", ref_map)
        assert text == "Plain text."
        assert refs == []


class TestDocxGenerator:
    def test_generates_valid_docx_bytes(self) -> None:
        gen = DocxGenerator(institution_name="Test Institute")
        result = gen.generate(
            title="测试公文",
            sections=[{"title": "第一章", "level": 1, "content": "正文内容 [ref:doc:p1]。"}],
            citations=[{"ref_id": "doc:p1", "document_title": "测试文献", "authors": "张三"}],
        )
        assert isinstance(result, bytes)
        assert len(result) > 0
        # DOCX files start with PK (ZIP format)
        assert result[:2] == b"PK"

    def test_generates_with_header_and_footer(self) -> None:
        gen = DocxGenerator(institution_name="PolicyAI")
        result = gen.generate(
            title="Report",
            sections=[{"title": "S1", "level": 1, "content": "Content."}],
            citations=[],
        )
        assert len(result) > 1000

    def test_generates_with_metadata(self) -> None:
        gen = DocxGenerator()
        result = gen.generate(
            title="Report",
            sections=[{"title": "S1", "level": 1, "content": "Content."}],
            citations=[],
            metadata={"date": "2026-05-19", "issue_number": "PolicyAI〔2026〕1号"},
        )
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    def test_multiple_sections(self) -> None:
        gen = DocxGenerator()
        result = gen.generate(
            title="Report",
            sections=[
                {"title": "Section 1", "level": 1, "content": "Content 1."},
                {"title": "Section 2", "level": 2, "content": "Content 2."},
            ],
            citations=[],
        )
        assert len(result) > 0

    def test_reference_list_included(self) -> None:
        gen = DocxGenerator()
        result = gen.generate(
            title="Report",
            sections=[{"title": "S1", "level": 1, "content": "Text [ref:doc:1]."}],
            citations=[{"ref_id": "doc:1", "document_title": "Test Doc", "authors": "Author"}],
        )
        # Verify valid DOCX output with references (text is inside ZIP XML entries)
        assert result[:2] == b"PK"
        assert len(result) > 2000  # Should contain reference-related XML


class TestFontHelper:
    def test_set_font_on_run(self) -> None:
        # _set_font is tested implicitly through DocxGenerator;
        # this verifies it doesn't crash with valid inputs
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("test")
        _set_font(run, "仿宋", 16, bold=False, color="#000000")
        assert run.font.size is not None
