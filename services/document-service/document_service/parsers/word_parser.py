"""Word (.docx) parser using python-docx (M2-11).

Extracts full text, paragraph styles (heading/body), tables,
and embedded images with OCR recognition.
"""

from __future__ import annotations

import io
import logging
from contextlib import suppress
from typing import Any

from document_service.models import PageContent, ParsedContent, SectionInfo
from document_service.parsers.base import BaseParser
from document_service.parsers.image_extractor import extract_images_from_docx

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Parse .docx files using python-docx."""

    def supported_format(self) -> str:
        return "docx"

    def parse(self, file_data: bytes, filename: str) -> ParsedContent:
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError("python-docx is required for Word parsing") from e

        doc = Document(io.BytesIO(file_data))
        paragraphs: list[str] = []
        sections: list[SectionInfo] = []
        all_tables: list[dict[str, Any]] = []

        for para in doc.paragraphs:
            text = para.text
            style_name = para.style.name if para.style else ""

            # Detect headings from style
            if style_name.startswith("Heading") or style_name.startswith("heading"):
                level = 1
                with suppress(ValueError):
                    level = int(style_name.replace("Heading", "").replace("heading", "").strip())
                sections.append(SectionInfo(
                    title=text,
                    level=min(level, 6),
                    page_start=0,
                    start_char=sum(len(p) + 1 for p in paragraphs),
                ))

            paragraphs.append(text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            all_tables.append({"table_index": table_idx, "rows": rows})

        full_text = "\n\n".join(paragraphs)

        # ---- Extract and OCR embedded images ----
        ocr_images = extract_images_from_docx(file_data)
        if ocr_images:
            image_texts = []
            for img in ocr_images:
                ocr_text = img.get("ocr_text", "")
                if ocr_text:
                    image_texts.append(
                        f"[Image {img.get('image_index', 0)} OCR ({img.get('format', 'png')})]:\n{ocr_text}"
                    )
            if image_texts:
                full_text += "\n\n--- Embedded Images (OCR) ---\n\n" + "\n\n".join(image_texts)
            logger.info("DOCX: OCR'd %d embedded images for %s", len(ocr_images), filename)
        # ---- End image extraction ----

        return ParsedContent(
            full_text=full_text,
            pages=[PageContent(page_number=1, text=full_text, has_text_layer=True)],
            tables=all_tables,
            sections=sections,
            metadata_hints=self.extract_metadata_hints(file_data, filename),
            needs_ocr=False,
            ocr_images=ocr_images,
        )

    def extract_metadata_hints(self, file_data: bytes, filename: str) -> dict[str, Any]:
        """Extract Word document properties (M2-21)."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            props = doc.core_properties
            return {
                "title": props.title or "",
                "author": props.author or "",
                "date": str(props.created) if props.created else "",
                "source": props.last_modified_by or "",
                "word_metadata": {
                    "title": props.title,
                    "author": props.author,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                    "last_modified_by": props.last_modified_by,
                    "revision": props.revision,
                    "category": props.category,
                    "subject": props.subject,
                },
            }
        except Exception:
            return {}
